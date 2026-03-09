import docx
import os

file_path = r"d:\毕设\rqglsys\lunwen\文档存储地址\5自动化测试的实现.docx"
output_path = r"d:\毕设\rqglsys\lunwen\ch5_full.txt"

doc = docx.Document(file_path)

with open(output_path, "w", encoding="utf-8") as f:
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            f.write(f"PARA [{i}]: {text}\n")
            
    for i, table in enumerate(doc.tables):
        f.write(f"\nTABLE [{i}]:\n")
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", r"\n") for cell in row.cells]
            f.write(" | ".join(row_text) + "\n")

print(f"写入到 {output_path}")
