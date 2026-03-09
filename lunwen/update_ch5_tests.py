import docx
import os
import shutil

# 文件路径
docx_path = r"d:\毕设\rqglsys\lunwen\文档存储地址\5自动化测试的实现.docx"
test_code_dir = r"d:\毕设\rqglsys\lunwen\自动化测试代码"

if not os.path.exists(test_code_dir):
    os.makedirs(test_code_dir)

doc = docx.Document(docx_path)

# 表格对应的模块名
modules = [
    "test_login_auth",
    "test_cylinder_management",
    "test_gas_purchase",
    "test_order_dispatch"
]

# 核心代码替换内容
simplified_codes = [
    # 模块1简化
"""// 核心业务：边界验证与身份鉴权
// 前置条件：WebDriver已初始化并进入登录页

// 正向用例：通过合法凭据进行鉴权
input(username, "admin")
input(password, "123456")
click(login_btn)

// 核心断言：根据Title进行系统级状态判定
if "管理员后台" in current_title:
    assert_success()
else:
    assert_fail("鉴权链路阻断")

// 逆向用例：边界流测试验证拦截器
input(username, "admin")
input(password, "654321") // 错误密码
click(login_btn)

if "用户名或密码错误" in page_source:
    assert_success("越权请求被成功阻断")
""",
    # 模块2简化
"""// 核心业务：钢瓶实体的生命周期创建与持久化断言
// 前置条件：已通过公共鉴权模块 module_auth 登录系统，并进入钢瓶管理域

click(add_cylinder_btn)

// 数据注入模拟
bind_data(cylinderCode, "GP-2025-0099")
bind_data(factory, "金源压力容器有限公司")
// ... 其他元数据

// 持久化事务提交
click(save_cylinder_btn)
sleep(2) // 等待异步事务闭环

// 核心断言：通过页面DOM回显验证事务闭环
if target_cylinder_code in page_source:
    assert_success("实体持久化成功")
else:
    assert_fail("数据写入阻断")
""",
    # 模块3简化
"""// 核心业务：客户态下购气订单的组装与状态机校验
// 前置条件：已通过公共鉴权模块 module_auth(customer) 登录

click(buy_gas_btn)
select_spec("15kg")
input(buyNum, "5")

// 补全履约数据与凭据后提交
submit_order()

// 模拟支付流转确认
click(pay_btn)

// 核心断言：校验订单实体状态是否正确演化至挂起/待分配
if "订单提交成功" in page_source and "待分配" in page_source:
    assert_success("终端缔约成功，状态机初始化完毕")
else:
    assert_fail("流转链路断裂")
""",
    # 模块4简化
"""// 核心业务：管理态下的资源调度与任务边缘节点投递
// 前置条件：已通过公共鉴权模块 module_auth(admin) 登录

navigate_to("订单管理 -> 待分配")

// 捕获挂起态实体进行派单
select_pending_order()
click(assign_courier_btn)

// 将任务投递至空闲节点(配送员)
select_courier("courier01")
click(confirm_assign_btn)

// 核心断言：校验状态机演变及与调度资源的耦合一致性
if "已指派" in page_source and "配送员A" in page_source:
    assert_success("调度图谱校验：资源下发成功")
else:
    assert_fail("分布协同失败")
"""
]

# 备份原文件
backup_path = docx_path.replace(".docx", "_backup.docx")
shutil.copy2(docx_path, backup_path)

# 表格索引：表9，表10，表11，表12 (0-based is 9, 10, 11, 12, as per ch5_full.txt output)
# 在前面的分析中，表9是登陆，表10是钢瓶，表11是购气，表12是调度
for idx, tab_idx in enumerate([9, 10, 11, 12]):
    if tab_idx < len(doc.tables):
        table = doc.tables[tab_idx]
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            # 获取完整代码
            full_code = table.rows[0].cells[0].text
            
            # 写入到真实代码文件中，用中国程序员习惯的注释
            py_file = os.path.join(test_code_dir, f"{modules[idx]}.py")
            with open(py_file, "w", encoding="utf-8") as f:
                f.write("# -*- coding: utf-8 -*-\n")
                f.write(f"# 自动化测试模块: {modules[idx]}\n")
                f.write(full_code)
                
            # 修改Word文档内容为精简版
            table.rows[0].cells[0].text = simplified_codes[idx]

# 保存修改后的文档
doc.save(docx_path)
print("自动化测试代码已提取至真实文件夹并且Word文档已优化替换完成。")
