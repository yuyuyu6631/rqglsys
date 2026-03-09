# -*- coding: utf-8 -*-
import docx
import os
import glob
import sys

# 设置输出编码
sys.stdout.reconfigure(encoding='utf-8')

# 获取文档目录
doc_dir = r"D:\毕设\rqglsys\lunwen\文档存储地址"

# 获取最大的文件（完整论文）
files = glob.glob(os.path.join(doc_dir, "*.docx"))
files.sort(key=lambda x: os.path.getsize(x), reverse=True)
doc_path = files[0]
print(f"读取: {os.path.basename(doc_path)}")

doc = docx.Document(doc_path)
print(f"总表格数: {len(doc.tables)}")

# 打印所有表格的详细信息
print("\n=== 表格详细信息 ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- 表格 {i} ---")
    print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
    
    # 打印前几行的内容
    for row_idx in range(min(3, len(table.rows))):
        row_data = []
        for cell in table.rows[row_idx].cells:
            text = cell.text.strip().replace('\n', ' ')
            row_data.append(text[:30] if len(text) > 30 else text)
        print(f"  行{row_idx}: {row_data}")
