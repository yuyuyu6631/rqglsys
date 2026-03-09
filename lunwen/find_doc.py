# -*- coding: utf-8 -*-
import docx
import os
import glob

# 获取文档目录
doc_dir = r"D:\毕设\rqglsys\lunwen\文档存储地址"

# 列出所有docx文件
files = glob.glob(os.path.join(doc_dir, "*.docx"))
for f in files:
    print(f"Found: {f}")

# 尝试读取第一个文件
if files:
    doc_path = files[0]
    print(f"\n尝试读取: {os.path.basename(doc_path)}")
    try:
        doc = docx.Document(doc_path)
        print(f"成功! 段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")
    except Exception as e:
        print(f"错误: {e}")
