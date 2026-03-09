# -*- coding: utf-8 -*-
"""
插入截图到 Word 文档脚本
用于将第四章功能测试的截图插入到 "第四章 功能测试的设计和实现.docx"
"""
import sys
import os
import glob
# 设置 UTF-8 编码支持
if sys.platform.startswith('win'):
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import datetime
from pathlib import Path

# 路径配置 - 使用 Windows 原生路径
doc_path = r"D:\毕设\rqglsys\lunwen\文档存储地址\第四章 功能测试的设计和实现.docx"
screenshots_dir = r"D:\毕设\rqglsys\lunwen\第四章截图及用例"

def main():
    print("=" * 60)
    print("Chapter 4 Function Test Screenshot Insertion Tool")
    print("=" * 60)

    # 检查文档是否存在
    if not os.path.exists(doc_path):
        print("[Error] Document not found:", doc_path)
        return

    # 备份原文档
    now = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_path = doc_path.replace('.docx', '_backup_' + now + '.docx')
    shutil.copy2(doc_path, backup_path)
    print("\n[Info] Backup created:", backup_path)

    # 打开文档
    doc = Document(doc_path)
    print("[Info] Original document has", len(doc.paragraphs), "paragraphs")

    # 显示目录内容
    print("\n[Info] Directory contents:")
    dir_path = Path(screenshots_dir)
    image_files = sorted(dir_path.glob('*.png'))
    for f in image_files:
        print(f"  - {f.name}")

    # 定义每个图对应的图片和标题描述
    # 格式：(文件名模式前缀，标题，描述)
    screenshot_data = [
        # DL 模块 (1-3)
        ("图 4.1 DL", "[图 4.1] DL-001 - 正常流程多角色账号系统登录",
         "展示用户登录成功后根据不同角色跳转至不同页面的界面效果"),
        ("图 4.2 DL", "[图 4.2] DL-002 - 密码错误拦截提示",
         "展示输入错误密码时系统弹出红色高亮警告提示的界面效果"),
        ("图 4.3 DL", "[图 4.3] DL-003 - 未登录访问拦截",
         "展示未登录状态下尝试访问受保护路由被强制跳转回登录页的效果"),

        # GP 模块 (4-6)
        ("图 4.4 GP", "[图 4.4] GP-001 - 钢瓶档案列表页面",
         "展示钢瓶全生命周期管理模块中在库钢瓶列表的渲染效果"),
        ("图 4.5 GP", "[图 4.5] GP-002 - 数据看板统计预警",
         "展示系统对临近到期钢瓶数据的统计和报警显示效果"),
        ("图 4.6 GP", "[图 4.6] GP-003 - 新增钢瓶表单校验",
         "展示缺少必填字段时表单校验拦截的效果"),

        # GQ 模块 (7-9)
        ("图 4.7 GQ", "[图 4.7] GQ-001 - 规格智能阶梯计费",
         "展示选择不同规格时金额实时计算联动的效果"),
        ("图 4.8 GQ", "[图 4.8] GQ-002 - 配送地址缺失拦截",
         "展示清空配送地址时系统提示补充的地址必填验证效果"),
        ("图 4.9 GQ", "[图 4.9] GQ-003 - 下单成功过渡页",
         "展示订单创建成功后的绿色对勾过渡页效果"),

        # DD 模块 (10-12)
        ("图 4.10 DD", "[图 4.10] DD-001 - 待分配订单列表",
         "展示管理员查看客户提交的 Pending 状态订单列表效果"),
        ("图 4.11 DD", "[图 4.11] DD-002 - 配送员接单界面",
         "展示配送员端接收任务并标记完成的操作界面效果"),
        ("图 4.12 DD", "[图 4.12] DD-003 - 已分配订单防重复派发",
         "展示对非 Pending 状态订单分配请求被系统拦截的效果"),
    ]

    success_count = 0
    total = len(screenshot_data)

    for idx, (filename_prefix, caption_text, desc_text) in enumerate(screenshot_data, 1):
        # 查找匹配的图片文件（两种格式都可能存在）
        # 格式 1: 【图 4.x...】.png (中文括号)
        # 格式 2: 图 4.x... .png (简单格式)
        pattern1 = f"{screenshots_dir}\\{filename_prefix}*.png"
        pattern2 = f"{screenshots_dir}\\【{filename_prefix}*.png"

        matches = glob.glob(pattern1) + glob.glob(pattern2)

        if not matches:
            print(f"\n[{idx}/{total}] No image found matching: {filename_prefix}")
            continue

        img_path = matches[0]
        print(f"\n[{idx}/{total}] Processing: {os.path.basename(img_path)}")

        try:
            # 添加段落用于图片
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 插入图片 (限制宽度为 8 英寸)
            run = para.add_run()
            run.add_picture(img_path, width=Inches(8))

            # 添加图片标题
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption_text)
            caption_run.bold = True

            # 添加描述文字
            desc_para = doc.add_paragraph()
            desc_para.paragraph_format.space_before = 0.3
            desc_para.paragraph_format.space_after = 0.3
            desc_para.add_run(desc_text)

            print("      OK - Successfully inserted")
            success_count += 1
        except Exception as e:
            print(f"      FAILED - Error: {str(e)}")

    # 保存修改后的文档
    output_path = doc_path
    doc.save(output_path)

    print("\n" + "=" * 60)
    print("Processing complete! Inserted {} out of {} images".format(success_count, total))
    print("Output document:", output_path)
    print("Backup document:", backup_path)
    print("=" * 60)

if __name__ == "__main__":
    main()
