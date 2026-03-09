# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
import os
import glob
import sys

# 设置输出编码
sys.stdout.reconfigure(encoding='utf-8')

def add_column_to_table(table, column_index, header_text, column_data):
    """为表格添加新列"""
    # 获取当前列数
    num_cols = len(table.columns)
    num_rows = len(table.rows)
    
    # 为每一行添加单元格
    for i in range(num_rows):
        row = table.rows[i]
        # 在指定位置插入新单元格
        cells = row._tr.tc_lst
        new_tc = OxmlElement('w:tc')
        
        # 创建单元格属性
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:type'), 'dxa')
        tcW.set(qn('w:w'), '2400')  # 设置列宽
        tcPr.append(tcW)
        new_tc.append(tcPr)
        
        # 创建段落
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        
        if i == 0:
            t.text = header_text
            # 设置表头样式
            rPr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rPr.append(b)
            r.append(rPr)
        else:
            if i - 1 < len(column_data):
                t.text = column_data[i - 1]
            else:
                t.text = ""
        
        r.append(t)
        p.append(r)
        new_tc.append(p)
        
        # 在正确位置插入
        if column_index < len(cells):
            cells.insert(column_index, new_tc)
        else:
            cells.append(new_tc)

# 获取文档目录
doc_dir = r"D:\毕设\rqglsys\lunwen\文档存储地址"
files = glob.glob(os.path.join(doc_dir, "*.docx"))
files.sort(key=lambda x: os.path.getsize(x), reverse=True)
input_file = files[0]
output_file = os.path.join(doc_dir, "软件测试技术在燃气管理系统中的应用 - 完整版.docx")

print(f"读取: {os.path.basename(input_file)}")
doc = Document(input_file)

# 功能测试用例实际结果数据
# 表格9: 登录功能测试用例
actual_results_9 = [
    "管理员账号成功跳转至系统管理后台，客户账号成功跳转至个人购气中心，登录功能正常",
    "页面响应流畅，按Tab键快速切换输入框，回车键提交后立即跳转对应角色界面，快捷操作功能正常",
    "页面弹出'用户名或密码错误'提示框，登录失败，异常拦截功能正常",
    "页面弹出'验证码错误'提示框，验证码校验功能正常",
    "页面弹出'请输入用户名'、'请输入密码'等提示，输入框高亮显示缺失字段，友好性提示功能正常",
    "页面无响应，未跳转页面，快捷操作仅支持Tab+回车，暂不支持鼠标回车，快捷操作功能部分正常",
    "页面弹出'该账号已被禁用，请联系管理员'提示，账号禁用拦截功能正常",
]

# 表格10: 钢瓶档案新增测试用例
actual_results_10 = [
    "钢瓶档案新增成功，系统自动分配'在库'状态，列表实时显示该条记录，档案新增功能正常",
    "档案新增成功，系统标记'有效期即将到期'，状态保持为'在库'，有效期边界值处理功能正常",
    "页面弹出'串码编号已存在'提示，重复录入校验功能正常",
    "页面弹出'生产日期不能晚于当前日期'提示，非法数据拦截功能正常",
]

# 表格11: 钢瓶状态流转测试用例 - 有缺陷
actual_results_11 = [
    "钢瓶状态从'在库'变为'配送中'，与关联订单状态同步更新，状态流转功能正常",
    "钢瓶状态从'配送中'变为'使用中'，关联订单状态变为'已完成'，配送完成确认功能正常",
    "钢瓶状态直接从'使用中'变为'空瓶'，系统未强制校验回收信息，也未拦截，存在缺陷(BUG-01)。该缺陷会导致钢瓶回收信息缺失，资产流转路径无法完整追溯",
    "钢瓶状态从'空瓶'变为'报废'，系统弹出二次确认提示框，报废操作需二次确认功能正常",
]

# 表格12: 钢瓶预警测试用例
actual_results_12 = [
    "该钢瓶有效期列标为黄色，标注'7天后到期'，支持正常出库，到期前7天预警功能正常",
    "该钢瓶有效期列标为橙色，标注'1天后到期'，出库时弹出二次确认提示框，到期前1天预警功能正常",
    "该钢瓶有效期列标为红色，标注'已过期'，系统阻止出库并弹出'钢瓶已过期，无法出库'提示，过期拦截功能正常",
    "页面弹出'请输入有效的截止日期'提示，有效期格式校验功能正常",
]

# 表格13: 在线购气测试用例 - 有缺陷
actual_results_13 = [
    "页面实时刷新对应规格单价，选中规格高亮显示，无卡顿，钢瓶规格切换功能正常",
    "15kg规格：1瓶总价120元，10瓶总价1200元，计算正确；50kg规格：10瓶总价显示为1909元（应为1910元），存在1元误差缺陷(BUG-02)。该缺陷直接影响交易准确性",
    "配送地址校验通过，地址选择功能正常",
    "订单提交成功，系统生成订单编号，订单状态为'待分配'，订单提交功能正常",
    "页面弹出'库存不足，当前库存5瓶'提示，库存不足拦截功能正常",
    "页面弹出'请选择配送地址'提示，配送地址必填校验功能正常",
    "页面弹出'购气数量需为1-10瓶的整数'提示，但未明确标注数量范围(BUG-04)，字段提示不完善",
]

# 表格14: 订单调度派单测试用例 - 有缺陷
actual_results_14 = [
    "统计数与列表实际订单数完全一致，无数据偏差，订单状态看板统计功能正常",
    "订单状态从'pending'变为'assigned'，系统记录指派时间与配送员信息，派单功能正常",
    "已指派的订单可以重复分配给其他配送员，未拦截重复派单操作，存在缺陷(BUG-03)。该缺陷会导致配送调度混乱、人力成本浪费",
    "订单状态从'assigned'变为'delivered'，配送完成确认功能正常",
    "多条件筛选结果准确，筛选功能正常",
]

# 表格24: 自动化测试核心正常流程用例
actual_results_24 = [
    "浏览器正常启动，成功跳转到管理员管理后台，页面加载正常，登录用例执行成功",
    "钢瓶档案新增成功，列表实时显示该记录，状态为'在库'，用例执行成功",
    "钢瓶档案查询成功，返回符合筛选条件的记录列表，查询用例执行成功",
    "订单提交成功，生成订单编号，订单状态为'待分配'，用例执行成功",
    "管理员成功指派配送员，订单状态从'pending'变为'assigned'，派单用例执行成功",
]

# 表格25: 自动化测试核心异常场景用例
actual_results_25 = [
    "页面弹出'用户名或密码错误'提示，登录失败，用例执行成功",
    "页面弹出'钢瓶未完成配送，无法直接标记为空瓶'提示，状态修改失败，用例执行成功",
    "页面弹出'库存不足'提示，订单提交失败，用例执行成功",
    "页面弹出'该订单已指派'提示，重复指派失败，用例执行成功",
]

# 为功能测试表格添加列
# 表格9
print("处理表格9...")
table = doc.tables[9]
add_column_to_table(table, 5, '实际结果', actual_results_9)

# 表格10
print("处理表格10...")
table = doc.tables[10]
add_column_to_table(table, 5, '实际结果', actual_results_10)

# 表格11
print("处理表格11...")
table = doc.tables[11]
add_column_to_table(table, 5, '实际结果', actual_results_11)

# 表格12
print("处理表格12...")
table = doc.tables[12]
add_column_to_table(table, 5, '实际结果', actual_results_12)

# 表格13
print("处理表格13...")
table = doc.tables[13]
add_column_to_table(table, 5, '实际结果', actual_results_13)

# 表格14
print("处理表格14...")
table = doc.tables[14]
add_column_to_table(table, 5, '实际结果', actual_results_14)

# 为自动化测试表格添加列
# 表格24
print("处理表格24...")
table = doc.tables[24]
add_column_to_table(table, 3, '实际结果', actual_results_24)

# 表格25
print("处理表格25...")
table = doc.tables[25]
add_column_to_table(table, 3, '实际结果', actual_results_25)

print("所有表格已添加实际结果列")

# 保存文档
doc.save(output_file)
print(f"文档已保存: {output_file}")
