# -*- coding: utf-8 -*-
import docx
import os

doc_path = r"D:\毕设\rqglsys\lunwen\文档存储地址\软件测试技术在燃气管理系统中的应用.docx"
doc = docx.Document(doc_path)

print(f"总表格数: {len(doc.tables)}")
print(f"总段落数: {len(doc.paragraphs)}")

# 打印所有段落标题
print("\n=== 文档结构 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and (text.startswith("第") or text.startswith("表") or text.startswith("图") or text.startswith("4.") or text.startswith("5.") or text.startswith("6.") or text.startswith("2.")):
        print(f"[{i}] {text[:80]}")

# 打印表格结构
print("\n=== 表格结构 ===")
for i, table in enumerate(doc.tables):
    print(f"表格 {i}: {len(table.rows)} 行 x {len(table.columns)} 列")
    # 打印表头
    if len(table.rows) > 0:
        header = [cell.text[:15] for cell in table.rows[0].cells]
        print(f"  表头: {header}")
