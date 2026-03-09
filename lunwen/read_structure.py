# -*- coding: utf-8 -*-
import docx
import os
import glob
import sys

# 设置输出编码
sys.stdout.reconfigure(encoding='utf-8')

# 获取文档目录
doc_dir = r"D:\毕设\rqglsys\lunwen\文档存储地址"

# 获取所有docx文件及其大小
files = glob.glob(os.path.join(doc_dir, "*.docx"))
file_info = []
for f in files:
    size = os.path.getsize(f)
    file_info.append((size, f))
    
# 按大小排序，最大的应该是完整论文
file_info.sort(reverse=True)
for size, f in file_info:
    print(f"大小: {size/1024/1024:.2f} MB - {os.path.basename(f)}")

# 使用最大的文件
doc_path = file_info[0][1]
print(f"\n使用最大文件: {os.path.basename(doc_path)}")

try:
    doc = docx.Document(doc_path)
    print(f"成功! 段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")
    
    # 打印章节标题
    print("\n=== 文档章节结构 ===")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 2:
            # 检测章节标题
            if any(keyword in text for keyword in ["第", "章", "功能测试", "自动化测试", "性能测试", "系统介绍"]):
                print(f"[P{i}]: {text[:80]}")

except Exception as e:
    print(f"错误: {e}")
